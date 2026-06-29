from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("output/documents/wolffia_pipseq_data_generation_protocol.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_font(run, name="Calibri", size=11, color=None, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text="", style=None, bold_start=None):
    p = doc.add_paragraph(style=style)
    if bold_start and text.startswith(bold_start):
        r = p.add_run(bold_start)
        set_font(r, bold=True)
        r2 = p.add_run(text[len(bold_start):])
        set_font(r2)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_bullets(doc, items, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    for item in items:
        p = doc.add_paragraph(style=style)
        for idx, part in enumerate(item if isinstance(item, list) else [item]):
            r = p.add_run(part[0] if isinstance(part, tuple) else part)
            if isinstance(part, tuple):
                set_font(r, bold=part[1] == "bold", italic=part[1] == "italic")
            else:
                set_font(r)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_font(r)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_font(run, size=16 if level == 1 else 13 if level == 2 else 12,
                 color="2E74B5" if level in (1, 2) else "1F4D78", bold=True)
    return p


def add_note(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [6.2])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    r = p.add_run(f"{label}: ")
    set_font(r, bold=True, color="1F3A5F")
    r2 = p.add_run(text)
    set_font(r2)
    doc.add_paragraph()


def style_document(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.49)
    sec.footer_distance = Inches(0.49)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name in ["List Bullet", "List Bullet 2", "List Number"]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(11)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.1

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("Wolffia PIP-seq protocol")
    set_font(r, size=9, color="666666")


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "F2F4F7")
        p = cell.paragraphs[0]
        r = p.add_run(h)
        set_font(r, bold=True)
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            cell = row.cells[i]
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            r = p.add_run(value)
            set_font(r, size=10.5)
    doc.add_paragraph()


def main():
    doc = Document()
    style_document(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run("PIP-seq-Adapted Protocol for Generating New Wolffia Transcriptomic Data")
    set_font(r, size=22, color="0B2545", bold=True)

    subtitle = doc.add_paragraph()
    r = subtitle.add_run("Stepwise pilot-to-discovery protocol for Wolffia australiana single-cell or single-nucleus transcriptomics")
    set_font(r, size=12, color="555555", italic=True)

    meta = doc.add_paragraph()
    r = meta.add_run(f"Last updated: {date.today().strftime('%m/%d/%Y')} | Platform assumption: PIP-seq")
    set_font(r, size=10, color="555555")

    add_heading(doc, "Purpose", 1)
    add_para(doc, "This protocol is designed to generate a first Wolffia dataset that is technically clean enough to recover broad transcriptional programs, distinguish real biological structure from preparation artifacts, and support a first Wolffia-native atlas analysis.")
    add_para(doc, "The central operational question remains upstream sample quality: which route produces the cleaner PIP-seq input for Wolffia, intact cells or nuclei? Because the project is now using PIP-seq, the protocol adds PIP-seq-specific gates for suspension quality, loading format, emulsification compatibility, and post-sequencing diagnostics.")

    add_heading(doc, "What Changes Because This Is PIP-seq", 1)
    add_bullets(doc, [
        "PIP-seq is microfluidics-free. It uses particle-templated emulsification to partition cells, barcoded hydrogel templates, and lysis reagents in uniform droplets using a standard vortexer rather than a Chromium-style microfluidic chip.",
        "The workflow is scalable in sample number and cell number. The Nature Biotechnology paper reports formats ranging from small tubes to 15 ml and 50 ml conical tubes, with emulsification completed in minutes. For this Wolffia project, scale should follow input quality, not the other way around.",
        "The critical sample deliverable is a clean single-particle suspension. Clumps, wall fragments, starch/plastid debris, residual enzymes, viscous carryover, or harsh lysis before encapsulation can reduce usable single-cell or single-nucleus transcriptomes.",
        "PIP-seq can support flexible processing formats, but the exact accepted input type, concentration, buffer, tube format, and loading target must be confirmed with the PIP-seq kit documentation, vendor, or sequencing core before the first run.",
        "Because PIP-seq captures poly(A) RNA on barcoded poly(T) hydrogel templates after partitioning and heat-triggered lysis, RNA preservation and rapid transition from preparation to compatible encapsulation conditions are more important than maximizing crude yield."
    ])

    add_note(doc, "Decision rule", "Do not use the largest PIP-seq loading format until Wolffia dissociation or nuclei isolation produces a low-clump, low-debris suspension that passes microscopy and counting QC.")

    add_heading(doc, "Scope", 1)
    add_bullets(doc, [
        "Wolffia culture standardization and pre-harvest QC",
        "Matched pilot comparison of intact-cell and nuclei preparation, if the PIP-seq provider supports both input types",
        "PIP-seq input acceptance criteria and loading-plan selection",
        "Execution of the first full Wolffia discovery run using one locked upstream route",
        "Immediate PIP-seq-aware computational QC after sequencing"
    ])
    add_para(doc, "This protocol does not claim that one Wolffia dissociation recipe or PIP-seq loading format is universally optimal. The published Wolffia protoplast workflow remains a useful starting benchmark, while the PIP-seq paper supplies the downstream platform logic: fast particle-templated emulsification, flexible container formats, and barcoded single-cell transcriptome capture without specialized microfluidic hardware.")

    add_heading(doc, "Recommended Baseline Biological Design", 1)
    add_bullets(doc, [
        "One Wolffia line or stock",
        "One baseline vegetative growth condition",
        "Three biological replicates for the first real run",
        "One fixed harvest window within the light cycle",
        "One locked preparation route after the pilot: intact cells or nuclei",
        "One locked PIP-seq loading plan and library-prep batch strategy"
    ])
    add_para(doc, "If the line is already maintained successfully in the lab, keep the maintenance environment unchanged during the pilot phase unless there is a strong reason to switch. The most important rule is consistency, not unnecessary re-optimization.")

    add_heading(doc, "Recommended Starting Culture Environment", 1)
    add_para(doc, "If a local standard already exists, use that standard. If no local standard exists, use this benchmark baseline adapted from the Wolffia australiana workflow previously identified for this project:")
    add_bullets(doc, [
        "Axenic liquid culture, if possible",
        "0.5x Schenk and Hildebrandt medium with 0.1% sucrose",
        "Medium adjusted to pH 6.7",
        "12-hour light / 12-hour dark photoperiod",
        "24 degrees C growth temperature",
        "Approximately 100 uE light intensity",
        "Weekly subculture by transferring about 10 fronds",
        "Harvest at the same time of day for all matched replicates",
        "For the published benchmark-style protoplast route, plan roughly 200 to 300 plants per pilot sample and avoid old or stressed material"
    ])

    add_heading(doc, "Required Materials and Equipment", 1)
    add_heading(doc, "Culture and Sampling", 2)
    add_bullets(doc, ["Healthy Wolffia stock", "Growth medium and sterile culture vessels", "Sterile pipettes, tips, forceps, and sampling tools", "Metadata sheet or electronic sample log", "Brightfield microscope or stereomicroscope and imaging device"])
    add_heading(doc, "Intact-Cell Pilot", 2)
    add_bullets(doc, ["Fresh sterile razor blade or fine disruption tool", "Osmotic stabilization buffer", "Plant cell-wall digestion reagents", "100 um sieve or mesh filter and 40 um cell strainer", "Low-speed centrifuge", "Hemocytometer or automated counter", "Optional viability dye"])
    add_heading(doc, "Nuclei Pilot", 2)
    add_bullets(doc, ["Ice bucket or cold block", "Nuclei isolation buffer", "Gentle homogenization setup such as a Dounce homogenizer", "Mesh filters or strainers", "Refrigerated centrifuge", "Nuclei counting setup", "Optional DNA stain"])
    add_heading(doc, "PIP-seq-Specific", 2)
    add_bullets(doc, [
        "PIP-seq kit or core-provided reagents, including barcoded hydrogel templates, oil, lysis/capture reagents, reverse-transcription and amplification reagents, and library-prep reagents",
        "Tube or plate format approved for the selected scale, such as small tube, microwell plate, 15 ml conical, or 50 ml conical format",
        "Standard vortexer and adapter compatible with the selected container format",
        "Thermal cycler or heating system compatible with PIP-seq heat-triggered lysis and downstream reactions",
        "Low-bind RNase-free tubes and plastics",
        "Library quantification and fragment analysis access",
        "Written PIP-seq provider specifications for input type, input buffer, allowable debris, concentration, loading target, and cleanup expectations"
    ])

    add_heading(doc, "Metadata That Must Be Recorded", 1)
    add_bullets(doc, [
        "Sample ID, biological replicate ID, stock or genotype ID",
        "Medium, temperature, photoperiod, culture age or transfer date, harvest date and time",
        "Operator, preparation route, preparation condition ID, harvest-to-processing time",
        "Cell or nuclei concentration, viability or nuclear-integrity observation, debris and clump score",
        "PIP-seq input volume, loading target, container format, barcoded template lot, reagent lot, vortex settings and duration, incubation timing, library batch, and sequencing batch",
        "Any cleanup, filtering, dilution, enrichment, or loading deviations"
    ])

    add_heading(doc, "Protocol", 1)
    add_heading(doc, "Step 1: Establish and Stabilize Wolffia Cultures", 2)
    add_numbered(doc, [
        "Start with one Wolffia line or stock only.",
        "Grow the material under one defined baseline vegetative condition.",
        "Maintain the same medium, light schedule, and temperature across all pilot cultures.",
        "Avoid mixing very old and very young cultures in the same experimental batch.",
        "Keep a culture log for at least several days before harvest."
    ])
    add_para(doc, "Output: healthy standardized cultures, stable maintenance conditions, and clear metadata records.")

    add_heading(doc, "Step 2: Perform Pre-Harvest Culture QC", 2)
    add_numbered(doc, [
        "Inspect representative fronds by brightfield microscopy or stereomicroscopy.",
        "Record visible morphology, frond size, budding status, and contamination.",
        "Estimate available biomass by frond count, wet weight, or both.",
        "Exclude cultures that are obviously stressed, contaminated, senescent, or highly heterogeneous."
    ])
    add_para(doc, "Go / no-go rule: proceed only if cultures are visibly healthy, reasonably uniform, free of obvious contamination, and sufficient in biomass for a split pilot.")

    add_heading(doc, "Step 3: Harvest Fresh Wolffia Material", 2)
    add_numbered(doc, [
        "Harvest pooled fronds or whole plants gently to avoid unnecessary mechanical damage.",
        "Use the same harvest method and timing for all matched pilot samples.",
        "Minimize time between harvest and the start of preparation.",
        "For nuclei preparations, place harvested material on ice immediately after harvest.",
        "For intact-cell preparations, move harvested material rapidly into the chosen wash or stabilization buffer."
    ])
    add_note(doc, "Critical record", "Record harvest-to-processing time for every sample. For PIP-seq, also record time from final suspension to emulsification or core handoff.")

    add_heading(doc, "Step 4: Split the Matched Pilot into Preparation Routes", 2)
    add_para(doc, "From the same starting biological material, divide the sample into Route A, intact-cell preparation, and Route B, nuclei preparation. The purpose of the pilot is technical comparison, not biological comparison.")
    add_para(doc, "If the PIP-seq provider does not support nuclei or has a separate nuclei-specific workflow, treat that information as a hard constraint and run only the compatible route in the first PIP-seq discovery run.")

    add_heading(doc, "Step 5A: Intact-Cell Preparation Pilot", 2)
    add_para(doc, "Goal: recover intact, viable single cells with minimal preparation-induced stress and minimal carryover into PIP-seq.")
    add_numbered(doc, [
        "Rinse freshly harvested Wolffia briefly to remove residual medium.",
        "For the published Wolffia benchmark route, collect about 200 to 300 plants and gently dice pooled material with a fresh sterile razor blade to improve enzyme access through the thick waterproof cuticle.",
        "Transfer diced material into a benchmark digestion solution containing 0.1 M KCl, 0.02 M MgCl2, 0.1% BSA, 0.08 M MES, 0.6 M mannitol, pH 5.5, plus 1.5% cellulase R-10, 1% maceroenzyme, and 0.5% pectolyase.",
        "Incubate under gentle agitation, keeping the shaker below 200 rpm, for a short pilot time course that does not exceed about 90 minutes unless deliberately testing a longer condition.",
        "Filter first through a 100 um sieve and then through a 40 um cell strainer.",
        "Wash recovered cells in the same buffer formulation without enzymes.",
        "Pellet and wash gently, for example near 1000 g at 20 degrees C followed by a later wash near 500 g at 20 degrees C if following the benchmark route.",
        "Resuspend in PIP-seq-compatible buffer or the provider-specified handoff buffer. Avoid residual enzyme mix, large debris, visible clumps, and viscous carryover.",
        "Count cells immediately, inspect morphology, and measure viability if supported."
    ])
    add_para(doc, "Pilot variables to compare: digestion time, enzyme strength, gentle agitation, optional pre-infiltration, and whole-frond handling versus gentle razor dicing.")

    add_heading(doc, "Step 5B: Nuclei Preparation Pilot", 2)
    add_para(doc, "Goal: recover clean, intact single nuclei while avoiding the stress burden of full wall digestion.")
    add_numbered(doc, [
        "Keep freshly harvested Wolffia cold from the start of processing.",
        "Transfer harvested material into cold nuclei isolation buffer.",
        "Homogenize gently just enough to release nuclei.",
        "Filter homogenate to remove large debris.",
        "Pellet and wash nuclei carefully under cold conditions.",
        "Count nuclei and inspect nuclear integrity.",
        "If appropriate, use a standard nuclei stain for counting or gating.",
        "Transfer nuclei into PIP-seq-compatible buffer only if the provider supports nuclei input."
    ])

    add_heading(doc, "Step 6: Apply the PIP-seq Input Acceptance Gate", 2)
    add_table(
        doc,
        ["Gate", "Accept", "Hold or redo"],
        [
            ("Single-particle state", "Mostly single cells or nuclei; limited aggregates after final filtration.", "Visible clumps, chains, or large plant fragments."),
            ("Debris burden", "Low debris by microscopy; no obvious wall fragments dominating the suspension.", "High background debris, plastid-rich carryover, or cloudy/viscous suspension."),
            ("Viability or integrity", "Cells look intact or nuclei are sharply defined and countable.", "Ruptured cells, lysed nuclei, or preparation-induced damage dominates."),
            ("Buffer compatibility", "Suspension is in provider-approved PIP-seq handoff buffer.", "Residual enzymes, detergents, high particulate load, or unknown additives remain."),
            ("Concentration/loading", "Within the provider's loading range for the selected format.", "Outside range, unverified concentration, or yield pursued at the expense of cleanliness."),
        ],
        [1.6, 2.45, 2.45],
    )

    add_heading(doc, "Step 7: Choose the Route and PIP-seq Scale for the Full Experiment", 2)
    add_para(doc, "Select the route using evidence, not preference. Choose the intact-cell route if viable cell recovery is reproducible, debris and clumping are manageable, and the preparation does not appear excessively harsh. Choose the nuclei route if cells are harsh, inconsistent, low-yield, or dominated by stress artifacts and the PIP-seq provider supports nuclei input.")
    add_para(doc, "Select the smallest scale that can answer the first biological question. PIP-seq can scale from small tube or plate-based workflows to conical-tube formats, but the first Wolffia discovery run should prioritize replicate consistency, clean suspensions, and interpretable QC over maximum cell count.")

    add_heading(doc, "Step 8: Run the PIP-seq Library Workflow", 2)
    add_para(doc, "Follow the current PIP-seq kit or core protocol exactly for reagent volumes, temperature steps, template input, oil addition, vortex settings, reverse transcription, amplification, tagmentation, indexing, and cleanup. The steps below define the Wolffia-specific control points rather than replacing the vendor protocol.")
    add_numbered(doc, [
        "Normalize each replicate to the same input type, buffer, concentration window, and loading target.",
        "Combine the Wolffia suspension with barcoded hydrogel templates and lysis/capture reagents according to the PIP-seq protocol.",
        "Add oil and emulsify using the approved vortex setting, duration, container format, and adapter. Record settings for every sample.",
        "Perform heat-triggered lysis and mRNA capture according to the PIP-seq workflow.",
        "Proceed through reverse transcription, whole-transcriptome amplification, tagmentation or fragmentation, indexing PCR, and cleanup according to the provider workflow.",
        "Quantify libraries and check fragment distributions before sequencing.",
        "Sequence biological replicates in a balanced way to avoid obvious batch imbalance."
    ])
    add_note(doc, "PIP-seq handling note", "For Wolffia, the main platform-specific risk is not the absence of microfluidics; it is introducing plant debris or aggregates into a droplet-partitioning workflow. A cleaner lower-yield input is usually preferable to a high-yield dirty input.")

    add_heading(doc, "Step 9: Perform Immediate Post-Sequencing QC", 2)
    add_bullets(doc, [
        "Reads per cell or nucleus and UMI counts per cell or nucleus",
        "Genes detected per cell or nucleus",
        "Barcode-rank profile and estimated recovered cell or nucleus count",
        "Probable doublet burden and unusually high-count barcodes",
        "Cross-sample or cross-replicate mixing if sample multiplexing was used",
        "Organellar fraction, especially chloroplast and mitochondrial signal",
        "Ambient-RNA-like background and plant debris-associated expression",
        "Stress-like artifact burden from dissociation or nuclei isolation",
        "Replicate concordance and stability of recovered clusters or neighborhoods"
    ])
    add_para(doc, "Minimum biological success criterion: the experiment is successful if it recovers broad expected programs reproducibly, such as proliferative or meristematic signal, photosynthetic or assimilation signal, transport or interface-associated signal, and developmental transition-like signal. Fine-grained cell labels are not required for the first run to be useful.")

    add_heading(doc, "Step 10: Decide on the Next Experiment", 2)
    add_para(doc, "Only after a clean baseline PIP-seq dataset exists should the project move to developmental contrasts, light or nutrient perturbations, osmotic or water-balance experiments, flowering or reproductive induction designs, or larger-scale PIP-seq loading formats.")

    add_heading(doc, "Recommended Timeline", 1)
    add_table(
        doc,
        ["Week", "Main work", "Decision point"],
        [
            ("1", "Standardize cultures, confirm PIP-seq provider requirements, finalize metadata tracking.", "Can the lab produce enough healthy material for a split pilot?"),
            ("2", "Run matched intact-cell and nuclei pilots if both are supported.", "Which input route gives the cleanest suspension?"),
            ("3", "Repeat the chosen route on fresh material if needed; confirm PIP-seq loading format.", "Has the input passed the PIP-seq acceptance gate?"),
            ("4", "Prepare PIP-seq libraries for three biological replicates and submit sequencing.", "Do libraries pass quantification and fragment QC?"),
            ("5 onward", "Run computational QC and first-pass biological interpretation.", "Is the baseline clean enough to expand the atlas design?"),
        ],
        [0.8, 3.1, 2.3],
    )

    add_heading(doc, "What Not to Do", 1)
    add_bullets(doc, [
        "Do not change culture conditions immediately before harvest.",
        "Do not harvest stressed or mixed-quality material.",
        "Do not force an intact-cell route when nuclei are clearly cleaner, or force nuclei if the PIP-seq provider does not support that input.",
        "Do not scale up to a large PIP-seq format before microscopy shows a clean single-particle suspension.",
        "Do not carry residual digestion enzymes, viscous material, or obvious plant debris into the PIP-seq reaction.",
        "Do not over-interpret dissociation stress, ambient RNA, or plant debris signatures as novel Wolffia biology."
    ])

    add_heading(doc, "Bottom Line", 1)
    add_para(doc, "The best first Wolffia PIP-seq experiment is the cleanest one: stable vegetative cultures, matched route testing where platform-compatible, evidence-based selection of cells or nuclei, strict PIP-seq input gates, three biological replicates, balanced library preparation, and immediate computational QC after sequencing.")

    add_heading(doc, "References", 1)
    refs = [
        "Clark I.C. et al. Microfluidics-free single-cell genomics with templated emulsification. Nature Biotechnology (2023). DOI: 10.1038/s41587-023-01685-z.",
        "Park J. et al. Genome of the world's smallest flowering plant, Wolffia australiana, helps explain its specialized physiology and unique morphology. Communications Biology (2021). DOI: 10.1038/s42003-021-02389-w.",
        "Hoang P.T.N. et al. The genome of Wolffia australiana facilitates discovery of genetic basis for aquatic adaptation in duckweeds. The Plant Cell (2022). DOI: 10.1093/plcell/koac068.",
        "Wu Y. et al. Streamlined spatial and environmental expression signatures characterize the minimalist duckweed Wolffia australiana. Genome Research 34(7):1106-1122 (2024). DOI: 10.1101/gr.279091.124.",
        "Public Wolffia reference datasets previously identified for benchmarking: PRJNA1124135 Wolffia scRNA-seq and PRJNA809022 Wolffia snRNA-seq."
    ]
    add_bullets(doc, refs)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    main()
